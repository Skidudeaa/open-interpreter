"""
Word Document Parser - Extract text and structure from .docx files.

Uses python-docx for parsing with support for:
- Paragraph and heading extraction
- Table extraction
- Metadata extraction
"""

from typing import Any

from .base import (
    BaseParser,
    DependencyMissingError,
    DocumentType,
    ParsedDocument,
)

# Lazy import for optional dependency
docx = None


def _get_docx():
    """Lazy load python-docx."""
    global docx
    if docx is None:
        try:
            import docx as _docx

            docx = _docx
        except ImportError:
            pass
    return docx


class DocxParser(BaseParser):
    """
    Parse Word documents (.docx) using python-docx.

    Features:
    - Paragraph extraction with style detection
    - Heading detection via styles
    - Table extraction
    - Metadata extraction
    """

    def can_parse(self, source: str) -> bool:
        """Check if source is a Word document."""
        return source.lower().endswith(".docx")

    def parse(self, source: str, **options: Any) -> ParsedDocument:
        """
        Parse a Word document (.docx).

        Args:
            source: Path to Word document
            **options: Additional options

        Returns:
            ParsedDocument with extracted text and structure

        Raises:
            DependencyMissingError: If python-docx is not installed
            FileNotFoundError: If file doesn't exist
        """
        docx_lib = _get_docx()
        if docx_lib is None:
            raise DependencyMissingError(
                "python-docx",
                "pip install 'open-interpreter[documents]' or pip install python-docx",
            )

        document = docx_lib.Document(source)

        text_parts = []
        sections = []
        tables = []
        parse_warnings = []

        current_section = None

        # Extract paragraphs
        for para in document.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            text_parts.append(para.text)

            # Detect headings by style
            style_name = para.style.name if para.style else ""

            if style_name.startswith("Heading"):
                # Extract heading level
                try:
                    level = int(style_name.replace("Heading", "").strip())
                except ValueError:
                    level = 1

                if current_section:
                    sections.append(current_section)

                current_section = {
                    "heading": text,
                    "level": level,
                    "content": "",
                }
            elif style_name == "Title":
                if current_section:
                    sections.append(current_section)
                current_section = {
                    "heading": text,
                    "level": 0,
                    "content": "",
                }
            elif current_section:
                current_section["content"] += text + "\n"

        # Add last section
        if current_section:
            sections.append(current_section)

        # Extract tables
        for table in document.tables:
            try:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                if table_data:
                    tables.append(table_data)
            except Exception as e:
                parse_warnings.append(f"Table extraction failed: {e}")

        # Extract metadata
        metadata = {}
        try:
            props = document.core_properties
            if props:
                metadata = {
                    "author": props.author,
                    "created": str(props.created) if props.created else None,
                    "modified": str(props.modified) if props.modified else None,
                    "last_modified_by": props.last_modified_by,
                    "subject": props.subject,
                    "keywords": props.keywords,
                    "category": props.category,
                    "comments": props.comments,
                }
                # Remove None values
                metadata = {k: v for k, v in metadata.items() if v is not None}
        except Exception as e:
            parse_warnings.append(f"Metadata extraction failed: {e}")

        # Get title from properties or first heading
        title = None
        try:
            if document.core_properties and document.core_properties.title:
                title = document.core_properties.title
        except Exception:
            pass
        if not title and sections:
            title = sections[0].get("heading")

        return ParsedDocument(
            text="\n\n".join(text_parts),
            document_type=DocumentType.DOCX,
            source=source,
            title=title,
            sections=sections,
            tables=tables,
            metadata=metadata,
            parse_warnings=parse_warnings,
        )
